from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("brain_tumor_ai_project_report.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def make_table(doc, headers, rows, widths=None, header_fill="1F4E79"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width
    return table


def add_diagram_box(doc, title, lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(31, 78, 121)
    for line in lines:
        rr = p.add_run(line + "\n")
        rr.font.name = "Consolas"
        rr.font.size = Pt(10)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Brain Tumor AI Project Report - Page ")
    add_page_number(footer)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("BRAIN TUMOR AI DESKTOP")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Thesis-Style Project Report on MRI Analysis, YOLO Inference, and AI Assisted Explanation")
    r.italic = True
    r.font.size = Pt(13)

    for line in [
        "Prepared for students seeking a complete academic project write-up",
        "Technology stack: React, Electron, FastAPI, Ultralytics YOLO, OpenCV, Pillow, and Gemini API",
        "Project type: Research and education prototype",
        "Disclaimer: Educational use only, not a medical diagnostic device",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "Abstract", 1)
    add_body(doc, (
        "This project presents a desktop-based brain tumor analysis system designed for educational and research use. "
        "The solution combines a React and Electron user interface with a Python FastAPI backend and a YOLO-based "
        "machine learning engine. Users can upload MRI images, receive a classification result, view an annotated "
        "image, and read an AI-generated explanation prepared through Gemini when available. The backend also "
        "exposes health and model inspection endpoints, supports batch analysis, and automatically searches for a "
        "working local port so the desktop shell can recover from conflicts during startup. The system is intentionally "
        "positioned as a prototype rather than a clinical device. Its main value lies in demonstrating how modern web "
        "technologies, computer vision models, and language-model explanations can be combined in a single workflow."
    ))
    add_body(doc, (
        "The report documents the problem statement, objectives, dataset preparation strategy, software architecture, "
        "module responsibilities, inference pipeline, evaluation approach, UI design, safety limitations, and future "
        "work. It is structured so that students can adapt it directly into a final-year project thesis."
    ))

    add_heading(doc, "1. Introduction", 1)
    add_body(doc, (
        "Brain tumors are among the most serious neurological conditions because they can affect cognition, movement, "
        "and life expectancy. Early interpretation of MRI scans is therefore a high-value task in medical imaging. "
        "However, real clinical diagnosis requires expert review, and many student projects fail when they only present "
        "a model without a usable system around it. This project addresses that gap by delivering a complete desktop "
        "application that allows image upload, inference, annotated visualization, status inspection, and friendly "
        "explanations in one package."
    ))
    add_body(doc, (
        "The application is built as a lightweight MVP that can be launched locally on Windows. The frontend is a React "
        "interface wrapped in Electron, while the backend exposes REST endpoints using FastAPI. The machine learning "
        "core uses Ultralytics YOLO and is designed to work with a trained brain tumor weight file placed in the models "
        "directory. In practice, this means the project can serve as both a software engineering demonstration and an "
        "AI-enabled academic prototype."
    ))

    add_heading(doc, "1.1 Problem Statement", 2)
    add_bullets(doc, [
        "Medical image analysis tools are often too technical for first-time users.",
        "Many student projects show a model result but do not provide a full workflow.",
        "Users need a visible pipeline from file upload to annotated output and explanation.",
        "Local startup failures, missing ports, and model-loading errors should be handled gracefully.",
    ])

    add_heading(doc, "1.2 Project Objectives", 2)
    add_numbered(doc, [
        "Develop a desktop application for MRI-based brain tumor analysis.",
        "Provide a simple upload-and-analyze workflow for individual and batch scans.",
        "Show model health, runtime device, and class information in the UI.",
        "Generate an annotated image and readable explanation for each analyzed scan.",
        "Support educational presentation of the result without claiming medical diagnosis.",
    ])

    add_heading(doc, "1.3 Scope of the Project", 2)
    add_body(doc, (
        "The scope of this project is limited to local desktop analysis of MRI images. The application focuses on a "
        "small but complete workflow: image upload, validation, inference, annotation, explanation, and result review. "
        "It does not attempt to handle hospital integration, electronic medical records, patient identity management, "
        "or multi-user collaboration. That limitation is intentional because it keeps the project suitable for a student "
        "thesis while still demonstrating a realistic end-to-end AI product."
    ))
    add_bullets(doc, [
        "Input: single MRI images or small batches of MRI images.",
        "Output: class label, confidence, annotated image, and explanation text.",
        "Environment: local Windows desktop and browser-based dev mode.",
        "Audience: students, supervisors, and research demonstrators.",
        "Not included: clinical certification, patient records, or live hospital deployment.",
    ])

    add_heading(doc, "2. System Overview", 1)
    add_body(doc, (
        "The architecture follows a three-layer pattern. The first layer is the user interface, which is implemented in "
        "React and packaged for desktop use through Electron. The second layer is the FastAPI backend, responsible for "
        "file validation, health checks, inference requests, annotated image serving, and explanation generation. The "
        "third layer is the model layer, which contains the trained YOLO weights and the image-processing logic that "
        "produces classification results or detection boxes."
    ))
    add_diagram_box(doc, "Figure 1. High-Level Architecture", [
        "[User]",
        "   |",
        "   v",
        "[React + Electron UI]",
        "   |  upload / analyze / status",
        "   v",
        "[FastAPI Backend]",
        "   |  validation / routing / file management",
        "   v",
        "[YOLO Engine + OpenCV]",
        "   |",
        "   +--> [Annotated Image Output]",
        "   +--> [Gemini Explanation Service]",
    ])

    add_heading(doc, "2.1 Runtime Flow", 2)
    add_numbered(doc, [
        "The user opens the desktop app or Vite-based UI.",
        "The frontend probes multiple API base URLs until a healthy backend is found.",
        "An MRI image is selected and uploaded from the analysis screen.",
        "The backend validates the file type and size before saving a temporary copy.",
        "The inference engine predicts the class and builds an annotated image.",
        "The explanation service optionally asks Gemini for a human-readable summary.",
        "Results are returned to the UI, displayed on screen, and stored in the recent queue.",
    ])

    add_heading(doc, "3. Technology Stack", 1)
    tech_rows = [
        ["React + Vite", "Frontend rendering, state management, file selection, and result display"],
        ["Electron", "Desktop shell for local deployment and startup orchestration"],
        ["FastAPI", "REST API layer for health, model info, analysis, and batch processing"],
        ["Ultralytics YOLO", "Classification and detection engine for MRI image analysis"],
        ["OpenCV", "Image loading, preprocessing, annotation, and bounding-box support"],
        ["Pillow", "Image handling and compatibility support"],
        ["Gemini API", "Natural-language explanation generation from model output"],
        ["Python", "Backend logic, training scripts, and evaluation tooling"],
    ]
    make_table(doc, ["Component", "Role"], tech_rows, widths=[Inches(1.6), Inches(4.9)])
    add_body(doc, (
        "The chosen stack is practical for a student project because the frontend can be iterated quickly, the backend "
        "remains lightweight, and the AI model can be swapped without changing the user interface contract. The codebase "
        "also keeps its responsibilities separated, which makes it easier to explain in a thesis."
    ))

    add_heading(doc, "3.1 Development Environment", 2)
    env_rows = [
        ["Operating system", "Windows desktop environment"],
        ["Backend runtime", "Python 3.10+ with FastAPI and Uvicorn"],
        ["Frontend runtime", "Node.js with Vite, React, and Electron"],
        ["AI runtime", "Ultralytics YOLO with OpenCV processing"],
        ["Auxiliary services", "Gemini API for explanation generation"],
    ]
    make_table(doc, ["Item", "Description"], env_rows, widths=[Inches(1.8), Inches(4.7)])

    add_heading(doc, "3.2 Non-Functional Goals", 2)
    add_bullets(doc, [
        "Keep the startup path resilient when the default backend port is occupied.",
        "Keep the interface simple enough for first-time users.",
        "Keep analysis results visible and easy to explain.",
        "Keep file handling temporary and predictable.",
        "Keep model selection replaceable without frontend rewrite.",
    ])

    add_heading(doc, "4. Dataset and Model Preparation", 1)
    add_body(doc, (
        "The training script expects an organized dataset directory with Train and Test splits. During preparation, it "
        "creates a working classification dataset with train, validation, and test folders. The split ratio is "
        "configurable, and the process copies images into a deterministic workspace so repeated runs remain comparable."
    ))
    add_bullets(doc, [
        "Training data is read from backend/dataset/Train.",
        "Testing data is read from backend/dataset/Test.",
        "Validation images are split from the training set using a seeded shuffle.",
        "Prepared data is written into backend/dataset/_prepared_classification.",
        "The base model defaults to yolov8n-cls.pt and the final model is stored as final.pt.",
    ])
    add_heading(doc, "4.1 Class Structure", 2)
    class_rows = [
        ["glioma", "Tumor class used by the project model"],
        ["meningioma", "Tumor class used by the project model"],
        ["notumor", "Negative class indicating no detected tumor"],
        ["pituitary", "Pituitary tumor class used by the project model"],
    ]
    make_table(doc, ["Class", "Meaning"], class_rows, widths=[Inches(1.8), Inches(4.7)])

    add_heading(doc, "4.2 Training Methodology", 2)
    add_body(doc, (
        "The training script converts the original dataset into a classification-friendly folder layout and then calls "
        "Ultralytics YOLO training. This is important because many datasets are stored as separate train and test "
        "folders, while the model expects a prepared workspace with train, validation, and test splits. The script uses "
        "a random seed for repeatable splitting, which is a good practice for academic experiments."
    ))
    add_numbered(doc, [
        "Read the training and test folders from the dataset directory.",
        "Shuffle the training images within each class using a fixed seed.",
        "Split a validation subset from the training images.",
        "Copy the prepared images into a temporary working dataset.",
        "Start YOLO training with the configured epoch count, image size, and batch size.",
        "Copy the best weights into a final model file for runtime inference.",
    ])

    add_heading(doc, "4.3 Why YOLO Classification Was Chosen", 2)
    add_body(doc, (
        "Although YOLO is widely known for object detection, this repository uses its classification pathway for the "
        "main workflow and keeps a fallback path for detection-style annotations. This design gives the project a very "
        "practical advantage: the backend can consume a single trained weight file and still return a structured result "
        "that the frontend can display in a consistent way. For a student project, this reduces integration complexity "
        "while keeping the model layer realistic and extensible."
    ))

    add_heading(doc, "5. Backend Architecture", 1)
    add_body(doc, (
        "The backend is the operational core of the application. It exposes /api/health to check readiness, /api/model "
        "to return model metadata, /api/analyze for single-image inference, /api/analyze/batch for multiple uploads, "
        "/api/explain for an extended analysis result, and /api/annotated/{filename} to stream the annotated image."
    ))
    add_bullets(doc, [
        "Upload validation blocks unsupported image extensions and large files.",
        "Temporary upload files are deleted after processing to avoid storage buildup.",
        "The backend auto-detects a usable local port when the default port is busy.",
        "Model loading uses candidate paths so the app can start with a locally available weight file.",
    ])
    add_heading(doc, "5.1 Inference Engine", 2)
    add_body(doc, (
        "The YOLOBrainTumorEngine class loads the trained model, reads the image through OpenCV, and runs prediction. "
        "If the model produces classification probabilities, the engine marks whether a tumor is detected, returns the "
        "highest-confidence class, and stores an annotated image. If the result falls back to a box-based output, the "
        "engine still produces structured detections with bounding boxes and confidence values."
    ))
    add_heading(doc, "5.2 Reported Output Fields", 2)
    output_rows = [
        ["tumor_detected", "Boolean status for quick interpretation"],
        ["predicted_type", "Selected class label"],
        ["best_confidence", "Top confidence score"],
        ["detections", "List of detected classes and scores"],
        ["annotated_image", "Filename of the annotated result"],
        ["processing_time_ms", "Measured inference time for the request"],
    ]
    make_table(doc, ["Field", "Purpose"], output_rows, widths=[Inches(1.9), Inches(4.6)])

    add_heading(doc, "5.3 API Endpoint Details", 2)
    api_rows = [
        ["/api/health", "Checks whether backend and model are ready"],
        ["/api/model", "Returns model metadata and class list"],
        ["/api/analyze", "Single-image analysis endpoint"],
        ["/api/analyze/batch", "Batch inference endpoint"],
        ["/api/explain", "Inference plus Gemini explanation"],
        ["/api/annotated/{filename}", "Streams annotated image output"],
    ]
    make_table(doc, ["Endpoint", "Purpose"], api_rows, widths=[Inches(2.0), Inches(4.5)])

    add_heading(doc, "5.4 Error Handling Strategy", 2)
    add_bullets(doc, [
        "Reject missing files and unsupported image extensions with HTTP errors.",
        "Reject files larger than the configured upload limit.",
        "Return a clear runtime error if the model file is not present.",
        "Delete temporary uploads after each inference run.",
        "Use safe fallback text when Gemini generation is unavailable.",
    ])

    add_heading(doc, "6. Frontend and User Experience", 1)
    add_body(doc, (
        "The frontend is designed as a dashboard rather than a plain upload form. This helps students present the project "
        "as a complete product. The navigation includes MRI Analysis, Batch Results, Model Status, and Settings. The "
        "analysis screen shows the uploaded image, the predicted status, confidence, detection list, annotated output, "
        "and a narrative explanation."
    ))
    add_bullets(doc, [
        "The app stores API settings locally so repeated launches remain stable.",
        "The UI probes backend health every few seconds to keep status current.",
        "The result queue allows users to revisit recent scans without re-uploading.",
        "The model screen shows readiness, runtime device, and current model path.",
        "The settings screen documents the educational and safety constraints.",
    ])
    add_heading(doc, "6.1 UI Section Summary", 2)
    ui_rows = [
        ["MRI Analysis", "Primary upload and inference workspace"],
        ["Batch Results", "History of recent processed scans"],
        ["Model Status", "Runtime health and model metadata"],
        ["Settings", "API base URL, Gemini model, and safety notes"],
    ]
    make_table(doc, ["Page", "Function"], ui_rows, widths=[Inches(1.6), Inches(4.9)])

    add_heading(doc, "6.2 User Interaction Flow", 2)
    add_numbered(doc, [
        "The user opens the application and waits for backend health to become online.",
        "The user switches to MRI Analysis and uploads a scan image.",
        "The user reviews the preview before starting inference.",
        "The user runs analysis and waits for the result cards to populate.",
        "The user inspects the predicted class, confidence, and annotated output.",
        "The user opens Batch Results or Model Status for additional project evidence.",
    ])

    add_heading(doc, "6.3 Frontend Design Rationale", 2)
    add_body(doc, (
        "The interface uses a dashboard layout so that the project looks more like a complete product and less like a "
        "single upload form. This is useful for a thesis defense because the reviewer can immediately see the system "
        "parts: input, status, analysis, and explanation. The side navigation also makes it easy to present the same "
        "screen structure during demos without losing context."
    ))

    add_heading(doc, "7. Explanation Layer", 1)
    add_body(doc, (
        "A useful thesis project should not stop at raw class output. This system includes a language-model explanation "
        "layer that converts the prediction into plain language. The service prepares a prompt describing the model "
        "result and asks Gemini to return JSON with a summary, confidence note, precautions, symptoms, escalation "
        "advice, and a disclaimer. If the API is unavailable, a safe fallback explanation is returned so the user still "
        "gets a complete result."
    ))
    add_diagram_box(doc, "Figure 2. Explanation Generation Flow", [
        "[Predicted class + confidence]",
        "            |",
        "            v",
        "[Prompt builder in gemini_service.py]",
        "            |",
        "            v",
        "[Gemini JSON response or fallback template]",
        "            |",
        "            v",
        "[Patient-friendly summary in the UI]",
    ])

    add_heading(doc, "7.1 Safety Content Logic", 2)
    add_body(doc, (
        "The explanation content intentionally avoids absolute diagnostic language. Instead, it frames the output as a "
        "screening result and reminds the user that a qualified doctor must review the scan. This is a good research "
        "decision because it prevents the project from appearing medically irresponsible. The fallback content also "
        "ensures the application remains usable when the external text generation service is offline."
    ))

    add_heading(doc, "8. Evaluation and Testing", 1)
    add_body(doc, (
        "The repository contains a GUI-based evaluation tool for the trained model. The tool loads the model and a YOLO "
        "dataset, runs batch predictions, calculates accuracy-style metrics, and renders a confusion matrix and class "
        "distribution charts. This is useful in a student thesis because it demonstrates that the model was tested, not "
        "just trained."
    ))
    add_bullets(doc, [
        "Top-1 accuracy is computed from predicted and target class labels.",
        "Precision, recall, and F1 are derived for a more complete performance view.",
        "Per-class accuracy helps identify class imbalance or weak categories.",
        "Confusion matrix visualization highlights where the model confuses classes.",
        "Sample prediction logs provide traceability for manual inspection.",
    ])
    add_heading(doc, "8.1 Suggested Academic Evaluation Table", 2)
    eval_rows = [
        ["Accuracy", "Top-1 class correctness on the held-out set"],
        ["Precision", "How many positive predictions are correct"],
        ["Recall", "How many true cases are recovered"],
        ["F1 Score", "Balance between precision and recall"],
        ["Inference Time", "Average response time per image"],
    ]
    make_table(doc, ["Metric", "Reason"], eval_rows, widths=[Inches(1.8), Inches(4.7)])

    add_heading(doc, "8.2 Test Scenarios", 2)
    test_rows = [
        ["Valid MRI image", "Inference should succeed and display a result"],
        ["Unsupported file type", "Upload should be rejected"],
        ["Oversized image", "Request should return a file size error"],
        ["Model missing", "UI should warn that best.pt is required"],
        ["Backend offline", "Frontend should show offline health status"],
        ["Gemini unavailable", "Fallback explanation should still appear"],
    ]
    make_table(doc, ["Scenario", "Expected outcome"], test_rows, widths=[Inches(2.2), Inches(4.3)])

    add_heading(doc, "8.3 Result Interpretation Guide", 2)
    add_bullets(doc, [
        "A higher confidence value means the model is more certain about the predicted class.",
        "Tumor detected does not mean the image confirms cancer clinically.",
        "The annotated image is a visualization aid, not a radiology report.",
        "Batch analysis is useful for demonstrating repeatability across multiple images.",
        "Performance should be discussed together with limitations and dataset quality.",
    ])

    add_heading(doc, "9. Security, Safety, and Limitations", 1)
    add_body(doc, (
        "Because this is a healthcare-related prototype, safety messaging is essential. The application repeatedly states "
        "that results are educational only and not a diagnosis. That design choice is appropriate and should remain in any "
        "future thesis or demo version. The backend also limits file types and file sizes, which reduces abuse and keeps "
        "processing predictable."
    ))
    add_bullets(doc, [
        "The app does not replace radiologist review.",
        "No clinical decision should be based on the output alone.",
        "The Gemini explanation is supportive text, not medical advice.",
        "The project depends on the presence of trained weights in the expected model path.",
        "Model accuracy will vary by dataset quality and class balance.",
    ])

    add_heading(doc, "9.1 Risk Analysis", 2)
    risk_rows = [
        ["Missing model file", "Prediction cannot start until weights are provided"],
        ["Poor image quality", "The model may produce low-confidence or wrong output"],
        ["Network outage", "Gemini explanation falls back to safe template text"],
        ["Port conflict", "Backend searches for a free local port automatically"],
        ["User misinterpretation", "Safety wording should be reviewed in demos"],
    ]
    make_table(doc, ["Risk", "Impact"], risk_rows, widths=[Inches(2.1), Inches(4.4)])

    add_heading(doc, "10. Future Improvements", 1)
    add_numbered(doc, [
        "Add a database layer for patient records and scan history.",
        "Introduce explainable AI overlays for saliency and attention visualization.",
        "Expand dataset handling with augmentation and cross-validation.",
        "Package the system into a more polished installer for non-technical users.",
        "Add role-based access control and audit logging for clinical environments.",
        "Replace placeholder fallback explanations with domain-approved templates.",
    ])

    add_heading(doc, "11. Conclusion", 1)
    add_body(doc, (
        "This Brain Tumor AI project demonstrates a full-stack approach to medical-image classification for academic use. "
        "It is stronger than a model-only submission because it includes the desktop interface, the backend API, the "
        "inference engine, explanation generation, status monitoring, and evaluation tooling. These pieces together make "
        "the system suitable as a thesis base for students who need a complete, structured, and explainable project."
    ))
    add_body(doc, (
        "From a software engineering perspective, the codebase shows modular design and practical error handling. From an "
        "AI perspective, it demonstrates how a trained YOLO model can be integrated into a usable product. From an "
        "academic perspective, the report captures the full workflow in a format that students can extend, defend, and "
        "present confidently."
    ))

    add_heading(doc, "Appendix A. Module Breakdown", 1)
    module_rows = [
        ["backend/main.py", "FastAPI routes, upload handling, health checks, and annotated file serving"],
        ["backend/inference.py", "YOLO loading, prediction, annotation, and bbox inference helpers"],
        ["backend/gemini_service.py", "Gemini prompt construction and fallback explanation text"],
        ["backend/train.py", "Dataset preparation and model training pipeline"],
        ["backend/evaluate_gui.py", "Desktop evaluation tool with charts and metrics"],
        ["frontend/src/main.jsx", "Primary user interface and navigation"],
        ["frontend/src/styles.css", "UI styling and layout system"],
    ]
    make_table(doc, ["Module", "Responsibility"], module_rows, widths=[Inches(2.0), Inches(4.5)])

    add_heading(doc, "Appendix B. Quick Thesis Defense Talking Points", 1)
    add_numbered(doc, [
        "What problem does the system solve?",
        "How does the backend validate and process MRI uploads?",
        "Why was YOLO selected for this project?",
        "How does the explanation layer improve usability?",
        "What are the system limitations and safety notes?",
        "How could the project be extended in future research?",
    ])

    add_heading(doc, "References", 1)
    add_bullets(doc, [
        "Ultralytics YOLO documentation and classification workflow.",
        "FastAPI official documentation for API design.",
        "OpenCV documentation for image preprocessing and annotation.",
        "Google Gemini API documentation for structured text generation.",
        "Python-docx documentation for report generation.",
    ])

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_document()
    print(f"Saved {OUT_PATH.resolve()}")
